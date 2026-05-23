import os
import re
import logging
from datetime import date, timedelta

_MONTHS_PT = {
    1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril",
    5: "maio", 6: "junho", 7: "julho", 8: "agosto",
    9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro",
}

def _date_pt(d: date) -> str:
    return f"{d.day} de {_MONTHS_PT[d.month]} de {d.year}"

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

OUTPUT_DIR   = "output"
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "template.docx")

COLOR_TEXT   = RGBColor(0x33, 0x33, 0x33)   # #333333 — template color
COLOR_RED    = RGBColor(0xFF, 0x00, 0x00)   # emendas deadline
COLOR_BLACK  = RGBColor(0x00, 0x00, 0x00)   # table label cells
COLOR_PURPLE = RGBColor(0x6B, 0x4E, 0x8B)   # objetivos box background
COLOR_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)   # objetivos box text


def _set_margins(doc: Document, top=2.0, bottom=2.0, left=3.0, right=1.5):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


def _set_default_font(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)


# ── Low-level helpers ────────────────────────────────────────────────────────

def _new_para(doc: Document, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Create a paragraph with zero space_before / space_after."""
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    return para


def _styled_run(para, text: str, *, bold=False, size=11, font_name: str | None = "Arial"):
    """Add a run with explicit color and optional bold/font/size."""
    r = para.add_run(text)
    r.bold           = bold
    r.font.size      = Pt(size)
    r.font.color.rgb = COLOR_TEXT
    if font_name:
        r.font.name = font_name
    return r


def _blank(doc: Document):
    """Empty paragraph — vertical spacer with zero spacing."""
    _new_para(doc)


def _add_divider(doc: Document):
    """Paragraph with bottom border in template color (#333333)."""
    para = _new_para(doc)
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "333333")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Prazos table ─────────────────────────────────────────────────────────────

def _add_prazos_table(doc: Document, pub_date: date) -> date:
    """3×2 deadline table positioned to the right side of the page.

      Eficácia:      DD/MM/YYYY a DD/MM/YYYY, prorrogável por mais 60 dias
      Sobrestamento: DD/MM/YYYY
      Emendas:       DD/MM/YYYY a DD/MM/YYYY   ← bold + red
    """
    eficacia_end  = pub_date + timedelta(days=59)   # day 60 (day 1 = publication)
    sobrestamento = pub_date + timedelta(days=45)
    emendas_end   = pub_date + timedelta(days=7)

    table = doc.add_table(rows=3, cols=2)
    tbl   = table._tbl

    # ── Table-level XML properties ────────────────────────────────────────────
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5081")
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), "3856")   # ≈ 6.8 cm indent → table on right side
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)

    tblBorders = OxmlElement("w:tblBorders")
    for bname in ("insideH", "insideV"):
        b = OxmlElement(f"w:{bname}")
        b.set(qn("w:val"),   "dotted")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    tblPr.append(tblBorders)

    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    # ── Column grid (widths) ──────────────────────────────────────────────────
    tblGrid = OxmlElement("w:tblGrid")
    for w in (1956, 3125):
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)
    tblPr_idx = list(tbl).index(tblPr)
    tbl.insert(tblPr_idx + 1, tblGrid)

    # ── Row data ──────────────────────────────────────────────────────────────
    rows_data = [
        (
            "Eficácia:",
            f"{pub_date.strftime('%d/%m/%Y')} a {eficacia_end.strftime('%d/%m/%Y')},"
            f" prorrogável por mais 60 dias",
            False, COLOR_BLACK,
        ),
        (
            "Sobrestamento:",
            sobrestamento.strftime("%d/%m/%Y"),
            False, COLOR_BLACK,
        ),
        (
            "Emendas:",
            f"{pub_date.strftime('%d/%m/%Y')} a {emendas_end.strftime('%d/%m/%Y')}",
            True, COLOR_RED,   # bold + red — action deadline
        ),
    ]

    for row_idx, (label, value, bold_val, val_color) in enumerate(rows_data):
        row = table.rows[row_idx]
        tr  = row._tr

        # Exact row height: 571 dxa
        trPr = tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = OxmlElement("w:trPr")
            tr.insert(0, trPr)
        trH = OxmlElement("w:trHeight")
        trH.set(qn("w:hRule"), "exact")
        trH.set(qn("w:val"),   "571")
        trPr.append(trH)

        # Cell widths
        for cell, width in zip(row.cells, (1956, 3125)):
            tc   = cell._tc
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"),    str(width))
            tcW.set(qn("w:type"), "dxa")
            tcPr.insert(0, tcW)

        # Label cell (col 0)
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p0.paragraph_format.space_after = Pt(0)
        rl = p0.add_run(label)
        rl.font.name      = "Arial"
        rl.font.size      = Pt(11)
        rl.font.color.rgb = COLOR_BLACK

        # Value cell (col 1)
        p1 = row.cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p1.paragraph_format.space_after = Pt(0)
        rv = p1.add_run(value)
        rv.font.name      = "Arial"
        rv.font.size      = Pt(11)
        rv.bold           = bold_val
        rv.font.color.rgb = val_color

    return emendas_end


def _add_atencao_box(doc: Document, emendas_end: date):
    """ATENÇÃO notice inside a single-cell bordered box."""
    table = doc.add_table(rows=1, cols=1)
    tbl   = table._tbl

    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Full-width table
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)

    # Black single border on all four sides
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "8")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        tblBorders.append(b)
    tblPr.append(tblBorders)

    cell = table.cell(0, 0)
    tc   = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", "60"), ("left", "108"), ("bottom", "60"), ("right", "108")):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"),    val)
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)

    # Line 1: "A T E N Ç Ã O !" in red bold centered
    p0 = cell.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after  = Pt(0)
    r0 = p0.add_run("A T E N Ç Ã O !")
    r0.bold           = True
    r0.font.size      = Pt(10)
    r0.font.color.rgb = COLOR_RED
    r0.font.name      = "Arial"

    # Line 2: main text with highlighted date
    p1 = cell.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after  = Pt(0)

    prefix = "AS EMENDAS DEVERÃO SER ENVIADAS PELO INFOLEG-AUTENTICADOR ATÉ 23h59min DO DIA "
    r1 = p1.add_run(prefix)
    r1.bold           = True
    r1.font.size      = Pt(10)
    r1.font.color.rgb = COLOR_BLACK
    r1.font.name      = "Arial"

    date_str = emendas_end.strftime("%d/%m/%Y") + "."
    r2 = p1.add_run(date_str)
    r2.bold           = True
    r2.font.size      = Pt(10)
    r2.font.color.rgb = COLOR_BLACK
    r2.font.name      = "Arial"
    # Yellow highlight on date
    rPr = r2._r.get_or_add_rPr()
    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), "yellow")
    rPr.append(highlight)


def _add_objetivos_box(doc: Document):
    """Purple filled paragraph with white bold centered text as section separator."""
    para = _new_para(doc, WD_ALIGN_PARAGRAPH.CENTER)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "6B4E8B")
    pPr.append(shd)
    r = para.add_run("OBJETIVOS DA MEDIDA PROVISÓRIA")
    r.bold           = True
    r.font.size      = Pt(12)
    r.font.color.rgb = COLOR_WHITE
    r.font.name      = "Arial"


# ── Document-level builders ──────────────────────────────────────────────────

def _set_header(doc: Document, title: str, subtitle: str):
    """Put title + subtitle in the Word page header (repeats on every page)."""
    header = doc.sections[0].header
    # Drop all existing paragraphs from the header body, then rebuild
    hdr_body = header._element
    for child in list(hdr_body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            hdr_body.remove(child)

    def _hdr_para(text: str) -> None:
        p = header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(text)
        r.bold           = True
        r.font.size      = Pt(12)
        r.font.name      = "Calibri"
        r.font.color.rgb = COLOR_TEXT

    _hdr_para(title)
    _hdr_para("")
    _hdr_para(subtitle)


def _add_metadata_line(doc: Document, label: str, value: str):
    para = _new_para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
    _styled_run(para, label + " ", bold=True)
    _styled_run(para, value,        bold=False)


def _add_section_heading(doc: Document, text: str):
    _blank(doc)
    para = _new_para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
    _styled_run(para, text, bold=True)
    _blank(doc)


def _add_labeled_block(doc: Document, label: str, body: str):
    _blank(doc)
    para = _new_para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
    _styled_run(para, label, bold=True)
    _blank(doc)
    _add_body_text(doc, body)


def _add_body_text(doc: Document, text: str, indent: bool = False):
    """Split on double newlines → one paragraph; single newline → line break."""
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    if not paragraphs:
        paragraphs = [text.strip()] if text.strip() else []
    for para_text in paragraphs:
        lines = para_text.splitlines()
        para = _new_para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
        if indent:
            para.paragraph_format.first_line_indent = Cm(1.25)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for j, line in enumerate(lines):
            r = para.add_run(line)
            r.font.size      = Pt(11)
            r.font.name      = "Arial"
            r.font.color.rgb = COLOR_TEXT
            if j < len(lines) - 1:
                r.add_break()


# ── Public API ───────────────────────────────────────────────────────────────

def write_nota_tecnica(mp: dict, content: dict, output_dir: str = OUTPUT_DIR) -> str:
    os.makedirs(output_dir, exist_ok=True)

    pub_date = date.fromisoformat(mp["data_publicacao"]) if mp.get("data_publicacao") else date.today()

    # Open template as base — inherits all styles, fonts and theme
    doc = Document(TEMPLATE_PATH) if os.path.exists(TEMPLATE_PATH) else Document()
    _set_margins(doc)
    _set_default_font(doc)

    # ── Prazos table — 4 blank lines below the page header ───────────────────
    for _ in range(4):
        _blank(doc)
    emendas_end = _add_prazos_table(doc, pub_date)
    _blank(doc)
    _add_atencao_box(doc, emendas_end)
    _blank(doc)

    # ── Title & subtitle → Word page header ──────────────────────────────────
    title    = content.get("titulo",    "NOTA TÉCNICA MP nº " + str(mp['numero']) + "/" + str(mp['ano']))
    subtitle = content.get("subtitulo", "Nota Informativa")
    _set_header(doc, title, subtitle)

    # ── Identification line ───────────────────────────────────────────────────
    para_ident = _new_para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
    para_ident.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    edicao = "Edição Extra do" if mp.get("edicao_extra") else "do"
    prefix_text = (
        f"A {edicao} Diário Oficial da União de {_date_pt(pub_date)}"
        f" publicou a Medida Provisória nº {mp['numero']}/{mp['ano']}, que "
    )
    r_prefix = para_ident.add_run(prefix_text)
    r_prefix.font.size      = Pt(11)
    r_prefix.font.name      = "Arial"
    r_prefix.font.color.rgb = COLOR_TEXT

    raw_ementa = mp.get("ementa", "")
    clean = re.sub(r"<[^>]+>", " ", raw_ementa)
    clean = re.sub(r"\s+", " ", clean).strip()
    # Truncate if the MP title reappears inside the ementa (Planalto page artefact)
    title_repeat = re.search(r"MEDIDA PROVIS[\xd3O]RIA\s+N", clean, re.IGNORECASE)
    if title_repeat:
        clean = clean[:title_repeat.start()].strip().rstrip(".,;")
    if not clean.endswith("."):
        clean = clean + "."
    ementa_text = "\u201c" + clean + "\u201d"
    r_ementa = para_ident.add_run(ementa_text)
    r_ementa.italic         = True
    r_ementa.font.size      = Pt(11)
    r_ementa.font.name      = "Arial"
    r_ementa.font.color.rgb = COLOR_TEXT


    _blank(doc)
    _add_objetivos_box(doc)
    _blank(doc)

    # ── Content: 5-paragraph structure (falls back to legacy resumo/alteracoes) ─
    if content.get("contexto"):
        for field in ("contexto", "dispositivos_centrais", "dispositivos_adicionais", "sintese", "fechamento"):
            text = content.get(field, "").strip()
            if text and text != " ":
                _add_body_text(doc, text, indent=True)
    else:
        if content.get("resumo"):
            _add_body_text(doc, content["resumo"], indent=True)
        if content.get("alteracoes"):
            _add_body_text(doc, content["alteracoes"], indent=True)

    # ── Fixed closing block (hardcoded — not generated by AI) ─────────────────
    _blank(doc)
    _add_divider(doc)
    _blank(doc)
    vigencia = _new_para(doc, WD_ALIGN_PARAGRAPH.CENTER)
    _styled_run(vigencia, "VIGÊNCIA: A Medida Provisória entra em vigor na data de sua publicação.", bold=True)
    _blank(doc)
    sig = _new_para(doc, WD_ALIGN_PARAGRAPH.CENTER)
    _styled_run(sig, "Assessoria da Liderança do Podemos", bold=True)

    # ── Save ──────────────────────────────────────────────────────────────────
    # ASCII-only filename: special chars (É, º) break GitHub Actions artifact ZIP
    filename = f"NOTA_TECNICA_-_MPV_n{mp['numero']}_de_{mp['ano']}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    logger.info("Nota técnica salva: %s", filepath)
    return filepath
